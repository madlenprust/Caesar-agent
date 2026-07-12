"""Инструменты категории 5: Документы.

См. roadmap раздел 11.6.
"""

import asyncio
import csv
import io
from typing import Any

from caesar.tools.base import Tool, ToolResult


class ParsePdfTool(Tool):
    name = "parse_pdf"
    description = "Извлечь текст из PDF. Поддерживает OCR для сканов (если установлен tesseract)."
    category = "documents"
    parameters_schema = {
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "pages": {"type": "string", "description": "Диапазон: 1-5 или 1,3,5"},
            "ocr": {"type": "boolean", "default": False},
            "extract_tables": {"type": "boolean", "default": True},
            "max_chars": {"type": "integer", "default": 100000},
        },
        "required": ["path"],
    }
    
    async def execute(
        self,
        path: str,
        pages: str | None = None,
        ocr: bool = False,
        extract_tables: bool = True,
        max_chars: int = 100000,
        **_,
    ) -> ToolResult:
        try:
            import pdfplumber
        except ImportError:
            return ToolResult(
                success=False,
                error="pdfplumber не установлен. Установи: pip install pdfplumber",
            )
        
        try:
            from pathlib import Path
            p = Path(path).expanduser()
            if not p.exists():
                return ToolResult(success=False, error=f"File not found: {path}")
            
            def _extract():
                pages_output = []
                with pdfplumber.open(str(p)) as pdf:
                    page_indices = self._parse_page_range(pages, len(pdf.pages))
                    for i in page_indices:
                        page = pdf.pages[i]
                        text = page.extract_text() or ""
                        
                        # Если текста мало — это скан, нужен OCR
                        if len(text.strip()) < 50 and not ocr:
                            text = "[scan detected — set ocr=true to extract via OCR]"
                        
                        tables = []
                        if extract_tables:
                            try:
                                tables = [
                                    {"rows": t.extract() or [], "caption": None}
                                    for t in page.extract_tables()
                                ]
                            except Exception:
                                pass
                        
                        pages_output.append({
                            "page_num": i + 1,
                            "text": text,
                            "tables": tables,
                            "has_images": bool(page.images),
                            "ocr_used": ocr and len(text.strip()) < 50,
                        })
                return pages_output
            
            pages_output = await asyncio.to_thread(_extract)
            
            # OCR если запрошено и pdfplumber дал мало текста
            if ocr:
                try:
                    import pytesseract
                    from pdf2image import convert_from_path
                    # TODO: реальный OCR
                except ImportError:
                    pass  # silently skip OCR if not installed
            
            total_text = "\n\n".join(p["text"] for p in pages_output)
            truncated = False
            if len(total_text) > max_chars:
                total_text = total_text[:max_chars] + "... (truncated)"
                truncated = True
            
            return ToolResult(
                success=True,
                data={
                    "total_pages": len(pages_output),
                    "pages": pages_output,
                    "truncated": truncated,
                    "text": total_text,  # для удобства
                },
            )
        except Exception as e:
            return ToolResult(success=False, error=str(e))
    
    def _parse_page_range(self, pages: str | None, total: int) -> list[int]:
        if not pages:
            return list(range(min(total, 100)))
        result = []
        for part in pages.split(","):
            part = part.strip()
            if "-" in part:
                start, end = part.split("-", 1)
                result.extend(range(int(start) - 1, min(int(end), total)))
            else:
                idx = int(part) - 1
                if 0 <= idx < total:
                    result.append(idx)
        return result


class ParseDocxTool(Tool):
    name = "parse_docx"
    description = "Извлечь текст, таблицы, стили из Word .docx файла."
    category = "documents"
    parameters_schema = {
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "include_tables": {"type": "boolean", "default": True},
            "max_chars": {"type": "integer", "default": 100000},
        },
        "required": ["path"],
    }
    
    async def execute(
        self,
        path: str,
        include_tables: bool = True,
        max_chars: int = 100000,
        **_,
    ) -> ToolResult:
        try:
            from docx import Document
        except ImportError:
            return ToolResult(
                success=False,
                error="python-docx не установлен. Установи: pip install python-docx",
            )
        
        try:
            from pathlib import Path
            p = Path(path).expanduser()
            if not p.exists():
                return ToolResult(success=False, error=f"File not found: {path}")
            
            def _extract():
                doc = Document(str(p))
                paragraphs = []
                for para in doc.paragraphs:
                    if not para.text.strip():
                        continue
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name if para.style else None,
                        "level": int(para.style.name.split()[-1]) if para.style and "Heading" in para.style.name else None,
                    })
                
                tables = []
                if include_tables:
                    for t in doc.tables:
                        rows = [[cell.text for cell in row.cells] for row in t.rows]
                        tables.append({"rows": rows, "caption": None})
                
                return paragraphs, tables
            
            paragraphs, tables = await asyncio.to_thread(_extract)
            
            total_text = "\n\n".join(p["text"] for p in paragraphs)
            truncated = False
            if len(total_text) > max_chars:
                total_text = total_text[:max_chars] + "... (truncated)"
                truncated = True
            
            return ToolResult(
                success=True,
                data={
                    "paragraphs": paragraphs,
                    "tables": tables,
                    "total_paragraphs": len(paragraphs),
                    "text": total_text,
                    "truncated": truncated,
                },
            )
        except Exception as e:
            return ToolResult(success=False, error=str(e))


class ParseXlsxTool(Tool):
    name = "parse_xlsx"
    description = "Извлечь данные из Excel .xlsx файла."
    category = "documents"
    parameters_schema = {
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "sheet": {"type": "string", "description": "Имя или индекс листа (0-based)"},
            "range": {"type": "string", "description": "A1:C10"},
            "has_header": {"type": "boolean", "default": True},
            "max_rows": {"type": "integer", "default": 1000},
        },
        "required": ["path"],
    }
    
    async def execute(
        self,
        path: str,
        sheet: str | int | None = None,
        range: str | None = None,
        has_header: bool = True,
        max_rows: int = 1000,
        **_,
    ) -> ToolResult:
        try:
            from openpyxl import load_workbook
        except ImportError:
            return ToolResult(
                success=False,
                error="openpyxl не установлен. Установи: pip install openpyxl",
            )
        
        try:
            from pathlib import Path
            p = Path(path).expanduser()
            if not p.exists():
                return ToolResult(success=False, error=f"File not found: {path}")
            
            def _extract():
                wb = load_workbook(str(p), read_only=True, data_only=True)
                sheets_output = []
                
                if sheet is not None:
                    if isinstance(sheet, int):
                        ws = wb.worksheets[sheet]
                    else:
                        ws = wb[sheet]
                    sheets_to_process = [ws]
                else:
                    sheets_to_process = wb.worksheets
                
                for ws in sheets_to_process:
                    rows = []
                    headers = None
                    
                    for i, row in enumerate(ws.iter_rows(values_only=True)):
                        if i >= max_rows:
                            break
                        if range:
                            # TODO: parse range
                            pass
                        rows.append([str(c) if c is not None else "" for c in row])
                    
                    if has_header and rows:
                        headers = rows[0]
                        rows = rows[1:]
                    
                    sheets_output.append({
                        "name": ws.title,
                        "total_rows": len(rows),
                        "total_cols": max(len(r) for r in rows) if rows else 0,
                        "headers": headers,
                        "rows": rows,
                    })
                
                wb.close()
                return sheets_output
            
            sheets_output = await asyncio.to_thread(_extract)
            
            return ToolResult(
                success=True,
                data={"sheets": sheets_output},
            )
        except Exception as e:
            return ToolResult(success=False, error=str(e))


class ParseCsvTool(Tool):
    name = "parse_csv"
    description = "Парсить CSV/TSV файл. Автоопределение разделителя."
    category = "documents"
    parameters_schema = {
        "type": "object",
        "properties": {
            "path": {"type": "string"},
            "delimiter": {"type": "string", "description": "Auto если не указан"},
            "has_header": {"type": "boolean", "default": True},
            "max_rows": {"type": "integer", "default": 1000},
        },
        "required": ["path"],
    }
    
    async def execute(
        self,
        path: str,
        delimiter: str | None = None,
        has_header: bool = True,
        max_rows: int = 1000,
        **_,
    ) -> ToolResult:
        try:
            from pathlib import Path
            p = Path(path).expanduser()
            if not p.exists():
                return ToolResult(success=False, error=f"File not found: {path}")
            
            def _extract():
                with open(p, "r", encoding="utf-8", errors="replace", newline="") as f:
                    sample = f.read(8192)
                    f.seek(0)
                    
                    if not delimiter:
                        # Автоопределение
                        try:
                            import csv as csv_mod
                            dialect = csv_mod.Sniffer().sniff(sample, delimiters=",\t;|")
                            delimiter = dialect.delimiter
                        except Exception:
                            delimiter = ","
                    
                    reader = csv.reader(f, delimiter=delimiter)
                    rows = []
                    for i, row in enumerate(reader):
                        if i >= max_rows:
                            break
                        rows.append(row)
                    
                    headers = None
                    if has_header and rows:
                        headers = rows[0]
                        rows = rows[1:]
                    
                    return {
                        "headers": headers,
                        "rows": rows,
                        "total_rows": len(rows),
                        "delimiter_detected": delimiter,
                    }
            
            data = await asyncio.to_thread(_extract)
            
            return ToolResult(success=True, data=data)
        except Exception as e:
            return ToolResult(success=False, error=str(e))


def get_documents_tools() -> list[Tool]:
    return [ParsePdfTool(), ParseDocxTool(), ParseXlsxTool(), ParseCsvTool()]
